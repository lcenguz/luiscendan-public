#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Informes APA desde Markdown
Convierte archivos .md a documentos Word (.docx) con formato APA 7ª edición
Basado en las normas APA y el estilo personal del usuario

IMPORTANTE: Lee archivos .docx directamente para extraer estilos reales
"""

import subprocess
import sys
import os
from pathlib import Path

# ============================================================================
# INSTALACIÓN AUTOMÁTICA DE DEPENDENCIAS
# ============================================================================

def instalar_dependencia(paquete, nombre_import=None):
    """Instala un paquete si no está disponible"""
    if nombre_import is None:
        nombre_import = paquete
    
    try:
        __import__(nombre_import)
        print(f"✅ {paquete} ya está lista.")
    except ImportError:
        print(f"📦 Instalando {paquete}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", paquete])
        print(f"✅ {paquete} instalada correctamente.")

# Instalar dependencias necesarias
instalar_dependencia("python-docx", "docx")

# ============================================================================
# IMPORTS
# ============================================================================

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import markdown
import re
from datetime import datetime

# ============================================================================
# ANALIZADOR DE ESTILOS DEL DOCUMENTO DE REFERENCIA
# ============================================================================

class AnalizadorEstilos:
    """Extrae estilos del documento .docx de referencia del usuario"""
    
    def __init__(self, archivo_referencia):
        """
        Args:
            archivo_referencia: Ruta al archivo .docx de referencia
        """
        self.archivo = Path(archivo_referencia)
        self.doc_ref = None
        self.estilos = {}
        
        if self.archivo.exists():
            self.doc_ref = Document(str(self.archivo))
            self.extraer_estilos()
    
    def extraer_estilos(self):
        """Extrae configuración de estilos del documento de referencia"""
        if not self.doc_ref:
            return
        
        # Extraer configuración de secciones (márgenes)
        if self.doc_ref.sections:
            section = self.doc_ref.sections[0]
            self.estilos['margenes'] = {
                'superior': section.top_margin,
                'inferior': section.bottom_margin,
                'izquierdo': section.left_margin,
                'derecho': section.right_margin
            }
        
        # Extraer estilos de párrafos
        for style in self.doc_ref.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                try:
                    self.estilos[style.name] = {
                        'fuente': style.font.name if style.font.name else 'Times New Roman',
                        'tamano': style.font.size if style.font.size else Pt(12),
                        'negrita': style.font.bold,
                        'cursiva': style.font.italic,
                        'interlineado': style.paragraph_format.line_spacing if style.paragraph_format.line_spacing else 2.0,
                        'alineacion': style.paragraph_format.alignment,
                        'sangria': style.paragraph_format.first_line_indent
                    }
                except:
                    pass
        
        print(f"📊 Estilos extraídos del documento de referencia: {len(self.estilos)} estilos")
    
    def obtener_estilo(self, nombre_estilo, default=None):
        """Obtiene un estilo específico"""
        return self.estilos.get(nombre_estilo, default)

# ============================================================================
# CONFIGURACIÓN APA
# ============================================================================

class ConfiguracionAPA:
    """Configuración de formato según APA 7ª edición"""
    
    # Fuente
    FUENTE = "Times New Roman"
    TAMANO_FUENTE = 12
    
    # Márgenes (en pulgadas)
    MARGEN = 1.0
    
    # Interlineado
    INTERLINEADO = 2.0  # Doble espacio
    
    # Sangría
    SANGRIA = 0.5  # Pulgadas
    
    # Títulos (jerarquía APA)
    NIVELES_TITULO = {
        1: {"negrita": True, "centrado": True, "mayusculas": True},
        2: {"negrita": True, "centrado": False, "mayusculas": False},
        3: {"negrita": True, "cursiva": False, "centrado": False, "mayusculas": False},
        4: {"negrita": True, "cursiva": True, "centrado": False, "mayusculas": False, "sangria": True},
        5: {"negrita": False, "cursiva": True, "centrado": False, "mayusculas": False, "sangria": True}
    }

# ============================================================================
# GENERADOR DE DOCUMENTOS APA
# ============================================================================

class GeneradorAPA:
    """Genera documentos Word con formato APA desde Markdown"""
    
    def __init__(self, archivo_md, metadatos=None, archivo_referencia=None):
        """
        Inicializa el generador
        
        Args:
            archivo_md: Ruta al archivo Markdown
            metadatos: Dict con título, autor, universidad, curso, profesor, fecha
            archivo_referencia: Ruta al .docx de referencia para extraer estilos (opcional)
        """
        self.archivo_md = Path(archivo_md)
        self.metadatos = metadatos or {}
        self.config = ConfiguracionAPA()
        
        # Si hay documento de referencia, extraer estilos
        self.analizador = None
        if archivo_referencia and Path(archivo_referencia).exists():
            print(f"📖 Usando estilos del documento de referencia: {archivo_referencia}")
            self.analizador = AnalizadorEstilos(archivo_referencia)
            self.doc = Document(str(archivo_referencia))
            # Limpiar contenido pero mantener estilos
            for element in self.doc.element.body:
                self.doc.element.body.remove(element)
        else:
            self.doc = Document()
        
    def configurar_documento(self):
        """Configura márgenes y estilos del documento"""
        # Configurar márgenes
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(self.config.MARGEN)
            section.bottom_margin = Inches(self.config.MARGEN)
            section.left_margin = Inches(self.config.MARGEN)
            section.right_margin = Inches(self.config.MARGEN)
        
        # Configurar estilo normal
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.config.FUENTE
        font.size = Pt(self.config.TAMANO_FUENTE)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.first_line_indent = Inches(self.config.SANGRIA)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    def crear_portada_estudiante(self):
        """Crea portada para estudiantes según APA 7ª"""
        # Título (centrado, negrita, mayúsculas)
        titulo = self.metadatos.get('titulo', 'TÍTULO DEL TRABAJO')
        p_titulo = self.doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(72)  # ~3 pulgadas desde arriba
        run = p_titulo.add_run(titulo.upper())
        run.font.name = self.config.FUENTE
        run.font.size = Pt(self.config.TAMANO_FUENTE)
        run.font.bold = True
        
        # Subtítulo (si existe)
        if 'subtitulo' in self.metadatos:
            p_subtitulo = self.doc.add_paragraph()
            p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_subtitulo.add_run(self.metadatos['subtitulo'])
            run.font.name = self.config.FUENTE
            run.font.size = Pt(self.config.TAMANO_FUENTE)
        
        # Espacio adicional
        self.doc.add_paragraph()
        
        # Autor
        autor = self.metadatos.get('autor', 'Nombre del Autor')
        p_autor = self.doc.add_paragraph()
        p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_autor.add_run(f"Autor: {autor}")
        run.font.name = self.config.FUENTE
        run.font.size = Pt(self.config.TAMANO_FUENTE)
        
        # Universidad
        universidad = self.metadatos.get('universidad', 'Universidad Alfonso X el Sabio')
        p_uni = self.doc.add_paragraph()
        p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_uni.add_run(universidad)
        run.font.name = self.config.FUENTE
        run.font.size = Pt(self.config.TAMANO_FUENTE)
        
        # Curso
        if 'curso' in self.metadatos:
            p_curso = self.doc.add_paragraph()
            p_curso.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_curso.add_run(self.metadatos['curso'])
            run.font.name = self.config.FUENTE
            run.font.size = Pt(self.config.TAMANO_FUENTE)
        
        # Profesor
        if 'profesor' in self.metadatos:
            p_prof = self.doc.add_paragraph()
            p_prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_prof.add_run(f"Profesor: {self.metadatos['profesor']}")
            run.font.name = self.config.FUENTE
            run.font.size = Pt(self.config.TAMANO_FUENTE)
        
        # Fecha
        fecha = self.metadatos.get('fecha', datetime.now().strftime("%d/%m/%Y"))
        p_fecha = self.doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_fecha.add_run(f"Fecha: {fecha}")
        run.font.name = self.config.FUENTE
        run.font.size = Pt(self.config.TAMANO_FUENTE)
        
        # Salto de página
        self.doc.add_page_break()
    
    def procesar_markdown(self):
        """Lee y procesa el archivo Markdown"""
        with open(self.archivo_md, 'r', encoding='utf-8') as f:
            contenido_md = f.read()
        
        # Procesar línea por línea para mantener formato
        lineas = contenido_md.split('\n')
        
        for linea in lineas:
            linea = linea.rstrip()
            
            # Títulos
            if linea.startswith('#'):
                self.agregar_titulo(linea)
            # Listas
            elif linea.startswith('- ') or linea.startswith('* '):
                self.agregar_lista(linea)
            # Párrafo vacío
            elif linea.strip() == '':
                continue
            # Párrafo normal
            else:
                self.agregar_parrafo(linea)
    
    def agregar_titulo(self, linea):
        """Agrega un título con formato APA"""
        # Contar nivel (número de #)
        nivel = len(linea) - len(linea.lstrip('#'))
        texto = linea.lstrip('#').strip()
        
        if nivel > 5:
            nivel = 5
        
        config_nivel = self.config.NIVELES_TITULO.get(nivel, {})
        
        p = self.doc.add_paragraph()
        
        # Alineación
        if config_nivel.get('centrado', False):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Sangría
        if config_nivel.get('sangria', False):
            p.paragraph_format.first_line_indent = Inches(self.config.SANGRIA)
        else:
            p.paragraph_format.first_line_indent = Inches(0)
        
        # Texto
        if config_nivel.get('mayusculas', False):
            texto = texto.upper()
        
        run = p.add_run(texto)
        run.font.name = self.config.FUENTE
        run.font.size = Pt(self.config.TAMANO_FUENTE)
        run.font.bold = config_nivel.get('negrita', False)
        run.font.italic = config_nivel.get('cursiva', False)
        
        # Espacio después del título
        p.paragraph_format.space_after = Pt(0)
    
    def agregar_parrafo(self, texto):
        """Agrega un párrafo con formato APA"""
        if not texto.strip():
            return
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(self.config.SANGRIA)
        
        # Procesar negritas y cursivas
        self.procesar_formato_inline(p, texto)
    
    def procesar_formato_inline(self, parrafo, texto):
        """Procesa formato inline (negritas, cursivas)"""
        # Patrón para detectar **negrita** y *cursiva*
        patron = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
        partes = re.split(patron, texto)
        
        for parte in partes:
            if not parte:
                continue
            
            run = parrafo.add_run()
            run.font.name = self.config.FUENTE
            run.font.size = Pt(self.config.TAMANO_FUENTE)
            
            if parte.startswith('**') and parte.endswith('**'):
                # Negrita
                run.text = parte[2:-2]
                run.font.bold = True
            elif parte.startswith('*') and parte.endswith('*'):
                # Cursiva
                run.text = parte[1:-1]
                run.font.italic = True
            elif parte.startswith('`') and parte.endswith('`'):
                # Código inline
                run.text = parte[1:-1]
                run.font.name = 'Courier New'
            else:
                # Texto normal
                run.text = parte
    
    def agregar_lista(self, linea):
        """Agrega un elemento de lista"""
        texto = linea.lstrip('- ').lstrip('* ').strip()
        
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        
        run = p.add_run(texto)
        run.font.name = self.config.FUENTE
        run.font.size = Pt(self.config.TAMANO_FUENTE)
    
    def generar(self, archivo_salida=None):
        """Genera el documento Word"""
        if archivo_salida is None:
            archivo_salida = self.archivo_md.with_suffix('.docx')
        
        print(f"📄 Generando documento APA desde: {self.archivo_md}")
        
        # Configurar documento
        self.configurar_documento()
        
        # Crear portada
        self.crear_portada_estudiante()
        
        # Procesar contenido Markdown
        self.procesar_markdown()
        
        # Guardar
        self.doc.save(archivo_salida)
        print(f"✅ Documento generado: {archivo_salida}")
        
        return archivo_salida

# ============================================================================
# FUNCIÓN PRINCIPAL
# ============================================================================

def convertir_md_a_apa(archivo_md, metadatos=None, archivo_salida=None, archivo_referencia=None):
    """
    Convierte un archivo Markdown a Word con formato APA
    
    Args:
        archivo_md: Ruta al archivo .md
        metadatos: Dict con título, autor, universidad, curso, profesor, fecha
        archivo_salida: Ruta del archivo .docx de salida (opcional)
        archivo_referencia: Ruta al .docx de referencia para extraer estilos (opcional)
    
    Returns:
        Ruta del archivo generado
    """
    generador = GeneradorAPA(archivo_md, metadatos, archivo_referencia)
    return generador.generar(archivo_salida)

# ============================================================================
# EJEMPLO DE USO
# ============================================================================

if __name__ == "__main__":
    # Ejemplo de uso
    metadatos_ejemplo = {
        'titulo': 'Plan de Ejecución: Caso D FinTech',
        'subtitulo': 'Detección de Fraude en Tiempo Real',
        'autor': 'José Luis Cendán Guzmán',
        'universidad': 'UNIVERSIDAD ALFONSO X EL SABIO',
        'curso': 'SM141503 - Plataformas Tecnológicas Cloud',
        'profesor': 'Nombre del Profesor',
        'fecha': '08/02/2026'
    }
    
    # Archivos
    archivo_md = r"C:\Users\joseluis.cendan\.gemini\antigravity\brain\804ff3dd-5807-4159-98d8-80026fd5f312\implementation_plan.md"
    archivo_referencia = r"c:\Github-Personal\luiscendan-private\estudios\master\1º_cuatrimestre\SM141502_Impacto_BigData_Negocios\Entregas\Entrega3\TrabajoFinal_JoseLuisCendanGuzman.docx"
    
    if os.path.exists(archivo_md):
        print("🚀 Generando documento con estilos de tu trabajo anterior...")
        convertir_md_a_apa(archivo_md, metadatos_ejemplo, archivo_referencia=archivo_referencia)
    else:
        print(f"❌ No se encontró el archivo: {archivo_md}")
        print("\n📝 Uso:")
        print("python md_to_apa_word.py")
        print("\nO desde Python:")
        print("from md_to_apa_word import convertir_md_a_apa")
        print("convertir_md_a_apa('mi_archivo.md', metadatos, archivo_referencia='tu_trabajo.docx')")
